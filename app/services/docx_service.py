import io
import re
from datetime import datetime
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MODULE_META = {
    'biz_definition': {'name': '사업 정의',  'icon': '①'},
    'env_analysis':   {'name': '환경 분석',  'icon': '②'},
    'value_design':   {'name': '가치 설계',  'icon': '③'},
    'revenue_model':  {'name': '수익 구조',  'icon': '④'},
    'execution':      {'name': '실행 체계',  'icon': '⑤'},
    'validation':     {'name': '검증·정제', 'icon': '⑥'},
}

ACCENT   = RGBColor(0x4f, 0x6e, 0xf7)
DARK     = RGBColor(0x1a, 0x1a, 0x2e)
GRAY     = RGBColor(0x55, 0x55, 0x55)
WHITE    = RGBColor(0xff, 0xff, 0xff)
WARN_BG  = RGBColor(0xff, 0xf8, 0xee)
INFO_BG  = RGBColor(0xee, 0xf2, 0xff)


# ── 헬퍼 ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'DDE3FF')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(13 if level == 1 else 11 if level == 2 else 10)
    return p


def _add_insight(doc, text, warn=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK
    # 왼쪽 테두리 흉내 — 실제 border는 XML로
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:color'), 'F5A623' if warn else '4F6EF7')
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def _add_implication(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK
    return p


# ── HTML → docx 파서 ──────────────────────────────────────────────

class _HtmlToDocx(HTMLParser):
    """AI 결과 HTML을 docx 요소로 변환하는 간단한 파서."""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self._cur_para = None
        self._cur_run  = None
        self._bold = False
        self._italic = False
        self._skip_tags = {'style', 'script'}
        self._skip = False

        # 테이블 상태
        self._in_table = False
        self._table = None
        self._table_row = None
        self._table_cell = None
        self._is_th = False
        self._col_idx = 0
        self._row_idx = 0
        self._col_count = 0          # 첫 행에서 결정
        self._pending_rows: list = []  # (is_header, [cell_texts])
        self._cur_cell_text = ''

        # 리스트
        self._in_list = False
        self._list_items: list = []
        self._cur_li_text = ''
        self._in_li = False

        # class 힌트
        self._cur_class = ''

    # ── 내부 유틸 ──

    def _flush_para(self):
        self._cur_para = None
        self._cur_run  = None

    def _ensure_para(self):
        if self._cur_para is None:
            self._cur_para = self.doc.add_paragraph()
            self._cur_para.paragraph_format.space_after = Pt(4)

    def _add_text(self, text):
        if not text:
            return
        if self._in_table:
            self._cur_cell_text += text
        elif self._in_li:
            self._cur_li_text += text
        else:
            self._ensure_para()
            run = self._cur_para.add_run(text)
            run.bold   = self._bold
            run.italic = self._italic
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK

    # ── 테이블 렌더링 ──

    def _render_table(self):
        if not self._pending_rows:
            return
        col_count = max(len(r[1]) for r in self._pending_rows) if self._pending_rows else 1
        table = self.doc.add_table(rows=0, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_border(table)

        for is_header, cells in self._pending_rows:
            row = table.add_row()
            for i, cell_text in enumerate(cells[:col_count]):
                cell = row.cells[i]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                run.font.size = Pt(9)
                if is_header:
                    run.bold = True
                    run.font.color.rgb = WHITE
                    _set_cell_bg(cell, '4F6EF7')
                else:
                    run.font.color.rgb = DARK

        self._pending_rows = []
        self.doc.add_paragraph()  # 테이블 이후 여백

    # ── 핸들러 ──

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        cls = attrs_dict.get('class', '')
        self._cur_class = cls

        if tag in self._skip_tags:
            self._skip = True
            return

        if tag in ('h3', 'h4'):
            self._flush_para()
            level = 2 if tag == 'h3' else 3
            self._cur_para = _add_heading(self.doc, '', level)
            self._bold = True

        elif tag == 'p':
            self._flush_para()
            if 'bcg-insight-warn' in cls:
                self._cur_para = self.doc.add_paragraph()
                self._cur_para.paragraph_format.left_indent = Cm(0.4)
                self._warn_insight = True
            elif 'bcg-insight' in cls:
                self._cur_para = self.doc.add_paragraph()
                self._cur_para.paragraph_format.left_indent = Cm(0.4)
                self._warn_insight = False
            else:
                self._cur_para = self.doc.add_paragraph()
            self._cur_para.paragraph_format.space_after = Pt(4)

        elif tag == 'div':
            if 'bcg-insight-warn' in cls:
                self._flush_para()
                self._cur_para = self.doc.add_paragraph()
                self._cur_para.paragraph_format.left_indent = Cm(0.4)
                pPr = self._cur_para._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), '12')
                left.set(qn('w:color'), 'F5A623')
                pBdr.append(left)
                pPr.append(pBdr)
            elif 'bcg-insight' in cls:
                self._flush_para()
                self._cur_para = self.doc.add_paragraph()
                self._cur_para.paragraph_format.left_indent = Cm(0.4)
                pPr = self._cur_para._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), '12')
                left.set(qn('w:color'), '4F6EF7')
                pBdr.append(left)
                pPr.append(pBdr)
            elif 'bcg-implication' in cls:
                self._flush_para()
                self._cur_para = self.doc.add_paragraph(style='List Bullet')
                self._cur_para.paragraph_format.space_after = Pt(2)

        elif tag == 'strong':
            self._bold = True

        elif tag == 'em':
            self._italic = True

        elif tag == 'br':
            if self._cur_para:
                self._cur_para.add_run('\n')

        elif tag == 'table':
            self._in_table = True
            self._pending_rows = []
            self._row_idx = 0
            self._flush_para()

        elif tag == 'tr':
            self._table_row = []
            self._col_idx = 0

        elif tag in ('th', 'td'):
            self._is_th = (tag == 'th')
            self._cur_cell_text = ''

        elif tag in ('ul', 'ol'):
            self._in_list = True
            self._list_items = []

        elif tag == 'li':
            self._in_li = True
            self._cur_li_text = ''

    def handle_endtag(self, tag):
        if tag in self._skip_tags:
            self._skip = False
            return

        if tag in ('h3', 'h4'):
            self._bold = False
            self._flush_para()

        elif tag == 'p':
            self._flush_para()

        elif tag == 'strong':
            self._bold = False

        elif tag == 'em':
            self._italic = False

        elif tag in ('th', 'td'):
            if self._table_row is not None:
                self._table_row.append(self._cur_cell_text)
                self._cur_cell_text = ''

        elif tag == 'tr':
            if self._table_row is not None:
                is_header = self._row_idx == 0
                self._pending_rows.append((is_header, self._table_row))
                self._table_row = None
                self._row_idx += 1

        elif tag == 'table':
            self._in_table = False
            self._render_table()

        elif tag in ('ul', 'ol'):
            for item in self._list_items:
                p = self.doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(item.strip())
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK
            self._in_list = False
            self._list_items = []

        elif tag == 'li':
            self._in_li = False
            self._list_items.append(self._cur_li_text)
            self._cur_li_text = ''

    def handle_data(self, data):
        if self._skip:
            return
        self._add_text(data)


def _html_to_docx(doc, html_text):
    if not html_text:
        return
    # 불필요한 태그 제거
    html_text = re.sub(r'<(span|a)[^>]*>', '', html_text)
    html_text = re.sub(r'</(span|a)>', '', html_text)
    parser = _HtmlToDocx(doc)
    parser.feed(html_text)


# ── 문서 생성 ──────────────────────────────────────────────────────

def _setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # 기본 스타일
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return doc


def _add_cover(doc, project, subtitle):
    now = datetime.now().strftime('%Y년 %m월 %d일')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run('사업모델캔버스')
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE
    run.bold = True
    # 배지 느낌 — 배경은 직접 XML
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '4F6EF7')
    rPr.append(shd)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(16)
    r2 = p2.add_run(project.name)
    r2.font.size = Pt(22)
    r2.bold = True
    r2.font.color.rgb = DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(6)
    r3 = p3.add_run(subtitle)
    r3.font.size = Pt(12)
    r3.font.color.rgb = GRAY

    doc.add_paragraph()

    meta_lines = []
    if project.organization:    meta_lines.append(f'수행 조직: {project.organization}')
    if project.business_manager: meta_lines.append(f'사업 담당: {project.business_manager}')
    if project.project_manager:  meta_lines.append(f'수행 담당: {project.project_manager}')
    meta_lines.append(f'출력일: {now}')

    for line in meta_lines:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_after = Pt(2)
        rm = pm.add_run(line)
        rm.font.size = Pt(9.5)
        rm.font.color.rgb = GRAY

    doc.add_page_break()


def generate_module_docx(project, module_type, ai_result) -> io.BytesIO:
    meta = MODULE_META.get(module_type, {'name': module_type, 'icon': ''})
    doc = _setup_doc()

    _add_cover(doc, project, meta['name'])

    _add_heading(doc, f"{meta['icon']} {meta['name']}", level=1)
    doc.add_paragraph()

    if ai_result:
        _html_to_docx(doc, ai_result)
    else:
        p = doc.add_paragraph('AI 분석 결과가 없습니다. 모듈에서 AI 분석을 실행해주세요.')
        p.runs[0].font.color.rgb = GRAY
        p.runs[0].italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_full_report_docx(project, modules, analyses) -> io.BytesIO:
    doc = _setup_doc()

    _add_cover(doc, project, '비즈니스 모델 설계 통합 보고서')

    # 목차
    _add_heading(doc, '목차', level=1)
    for i, module in enumerate(modules, 1):
        analysis = analyses.get(module['key'])
        if analysis and analysis.ai_result:
            status = '✓ AI 분석 완료'
        elif analysis and analysis.input_data:
            status = '- 데이터 입력됨'
        else:
            status = '○ 미작성'
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f'{i:02d}.  {module["name"]}')
        r.font.size = Pt(10)
        r.font.color.rgb = DARK
        rs = p.add_run(f'   {status}')
        rs.font.size = Pt(9)
        rs.font.color.rgb = GRAY

    doc.add_page_break()

    # 각 모듈 본문
    for module in modules:
        meta = MODULE_META.get(module['key'], {'name': module['key'], 'icon': ''})
        analysis = analyses.get(module['key'])
        ai_result = analysis.ai_result if analysis else None

        _add_heading(doc, f"{meta['icon']} {meta['name']}", level=1)
        doc.add_paragraph()

        if ai_result:
            _html_to_docx(doc, ai_result)
        else:
            p = doc.add_paragraph('AI 분석 결과가 없습니다.')
            p.runs[0].font.color.rgb = GRAY
            p.runs[0].italic = True

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
